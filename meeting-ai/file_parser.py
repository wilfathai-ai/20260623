"""Word(.docx) / PDF / テキストファイルからの本文抽出"""
from __future__ import annotations

from docx import Document
from pypdf import PdfReader


class FileParseError(Exception):
    """ファイル解析に失敗した際に投げる例外。メッセージはそのままユーザーに表示する想定。"""


def parse_docx(file_path: str) -> str:
    try:
        document = Document(file_path)
    except Exception as exc:
        raise FileParseError(f"Wordファイルの読み込みに失敗しました: {exc}") from exc

    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" / ".join(cells))

    text = "\n".join(paragraphs).strip()
    if not text:
        raise FileParseError("ファイルからテキストを抽出できませんでした。内容が空の可能性があります。")
    return text


def parse_pdf(file_path: str) -> str:
    try:
        reader = PdfReader(file_path)
    except Exception as exc:
        raise FileParseError(f"PDFファイルの読み込みに失敗しました: {exc}") from exc

    if reader.is_encrypted:
        raise FileParseError("暗号化されたPDFは読み込めません。")

    pages = []
    for page in reader.pages:
        page_text = page.extract_text() or ""
        if page_text.strip():
            pages.append(page_text.strip())

    text = "\n".join(pages).strip()
    if not text:
        raise FileParseError("PDFからテキストを抽出できませんでした。スキャン画像のみのPDFには対応していません。")
    return text


def parse_txt(file_path: str) -> str:
    encodings = ("utf-8", "shift_jis", "cp932")
    last_error: Exception | None = None
    for encoding in encodings:
        try:
            with open(file_path, "r", encoding=encoding) as f:
                text = f.read().strip()
            if not text:
                raise FileParseError("ファイルが空です。")
            return text
        except FileParseError:
            raise
        except Exception as exc:
            last_error = exc
            continue
    raise FileParseError(f"テキストファイルの読み込みに失敗しました: {last_error}")


def parse_file(file_path: str, filename: str) -> str:
    lower = filename.lower()
    if lower.endswith(".docx"):
        return parse_docx(file_path)
    if lower.endswith(".pdf"):
        return parse_pdf(file_path)
    if lower.endswith(".txt"):
        return parse_txt(file_path)
    raise FileParseError("対応していないファイル形式です。.docx / .pdf / .txt をアップロードしてください。")
